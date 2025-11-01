#!/usr/bin/env python3
"""
DOCX -> studies.json ingester for SoulStart.

- Detects study sections by headings like:
  "1.", "1)", "1 -", "Day 1", "Day 1:", etc.
- Extracts scripture lines ("Scripture:", "Verse:", "Text:")
- Aggregates outline paragraphs and bullet points into a normalized structure.

Usage:
  python tools/ingest.py
  python tools/ingest.py -i path/to/Week1_Devotion.docx -o devotions/studies.json
"""

from __future__ import annotations
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any

try:
    from docx import Document  # pip install python-docx
except ImportError as e:
    raise SystemExit("Missing dependency: python-docx. Install with: pip install python-docx") from e


# --- Paths (project-root aware) ---
BASE_DIR = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = BASE_DIR / "Week1_Devotion.docx"
DEFAULT_OUT  = BASE_DIR / "devotions" / "studies.json"

# --- Heuristics / regexes ---
# Start-of-study markers: "1.", "1)", "1 -", "Day 1", "Day 1:"
STUDY_START_RE = re.compile(
    r"^\s*(?:day\s*(\d+)\s*[:.\-]\s*|(\d+)\s*[\.\)\-:]\s*)(.*)$",
    flags=re.IGNORECASE
)

SCRIPTURE_RE = re.compile(r"^\s*(?:scripture|verse|text)\s*:\s*(.+)$", re.IGNORECASE)

# Bullets or numbered points within a study
BULLET_RE = re.compile(r"^\s*(?:[-*•–]|(\d+)[\.\)])\s+(.*)$")

# Cleaners
WS_RE = re.compile(r"\s+")


def _norm_ws(s: str) -> str:
    """Normalize whitespace and strip."""
    return WS_RE.sub(" ", s or "").strip()


def _is_heading(p_text: str) -> bool:
    """Does this paragraph start a new study?"""
    return bool(STUDY_START_RE.match(p_text))


def _maybe_scripture(p_text: str) -> str | None:
    """Extract scripture text if the paragraph looks like a Scripture line."""
    m = SCRIPTURE_RE.match(p_text)
    return _norm_ws(m.group(1)) if m else None


def parse_docx(docx_path: Path) -> List[Dict[str, Any]]:
    doc = Document(str(docx_path))
    studies: List[Dict[str, Any]] = []

    cur: Dict[str, Any] | None = None
    outline_parts: List[str] = []
    points: List[str] = []

    def push_current():
        if cur is None:
            return
        cur_out = _norm_ws(" ".join(outline_parts))
        cur_pts = [p for p in (x.strip() for x in points) if p]
        cur.setdefault("outline", cur_out)
        if "points" not in cur:
            cur["points"] = cur_pts
        else:
            # merge (shouldn’t happen, but safe)
            cur["points"].extend(cur_pts)
        # normalize keys presence
        cur.setdefault("title", "Study")
        cur.setdefault("scripture", "")
        cur.setdefault("resources", [])
        studies.append(cur.copy())

    for para in doc.paragraphs:
        raw = (para.text or "").strip()
        if not raw:
            continue

        text = _norm_ws(raw)

        # New study heading?
        if _is_heading(text):
            # flush previous
            push_current()
            outline_parts = []
            points = []

            m = STUDY_START_RE.match(text)
            # Extract the trailing title part (group 3) if present
            title_tail = _norm_ws(m.group(3) or "")
            cur = {
                "title": title_tail or f"Study {m.group(1) or m.group(2)}",
                "scripture": "",
                "outline": "",
                "points": [],
                "resources": []
            }
            continue

        if cur is None:
            # If the doc doesn't start with a heading, start one lazily.
            cur = {"title": "Study", "scripture": "", "outline": "", "points": [], "resources": []}
            outline_parts = []
            points = []

        # Scripture line?
        scr = _maybe_scripture(text)
        if scr:
            cur["scripture"] = scr
            continue

        # Bullet/numbered point?
        b = BULLET_RE.match(text)
        if b:
            # prefer the text portion after bullet
            pt = b.group(2) if b.group(2) else text
            points.append(_norm_ws(pt))
            continue

        # Fallback: add to outline blob
        outline_parts.append(text)

    # Flush last study
    push_current()

    return studies


def normalize_for_app(studies: List[Dict[str, Any]], limit: int | None = 6) -> List[Dict[str, Any]]:
    """Ensure the shape the app expects; optionally limit to first N studies."""
    if limit is not None:
        studies = studies[:limit]

    out: List[Dict[str, Any]] = []
    for s in studies:
        out.append({
            "title": s.get("title") or "Study",
            "scripture": s.get("scripture", ""),
            "outline": s.get("outline", ""),
            "points": s.get("points", []) or [],
            "resources": s.get("resources", []) or []
        })
    return out


def write_json(data: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"✅ Wrote {path} with {len(data) if isinstance(data, list) else 'N'} studies")


def main():
    ap = argparse.ArgumentParser(description="Ingest Week* DOCX into devotions/studies.json")
    ap.add_argument("-i", "--input", type=Path, default=DEFAULT_DOCX, help="Input DOCX file")
    ap.add_argument("-o", "--output", type=Path, default=DEFAULT_OUT, help="Output JSON file")
    ap.add_argument("--no-limit", action="store_true", help="Do not limit to first 6 studies")
    args = ap.parse_args()

    if not args.input.exists():
        raise SystemExit(f"❌ Input DOCX not found: {args.input}")

    studies = parse_docx(args.input)
    if not studies:
        raise SystemExit("❌ No studies detected. Check your headings (e.g., '1.' or 'Day 1: Title').")

    cleaned = normalize_for_app(studies, limit=None if args.no_limit else 6)

    # Simple console summary
    print("— Summary —")
    for i, s in enumerate(cleaned, 1):
        pts = len(s.get("points", []))
        print(f"{i}. {s.get('title','Study')} | scripture='{s.get('scripture','')}' | points={pts}")

    write_json(cleaned, args.output)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
DOCX → JSON + CSV converter for SoulStart Devotion.
Also provides a callable function `import_studies()` for Flask admin use.
"""

from __future__ import annotations
import re, json, csv
from pathlib import Path
from typing import List, Dict, Any
from docx import Document

BASE_DIR = Path(__file__).resolve().parents[1]
DEV_DIR  = BASE_DIR / "devotions"
DOCX_PATH = BASE_DIR / "Week1_Devotion.docx"
JSON_OUT  = DEV_DIR / "studies.json"
CSV_OUT   = DEV_DIR / "studies.csv"

STUDY_RE = re.compile(r"^\s*(?:day\s*(\d+)\s*[:.\-]\s*|(\d+)\s*[\.\)\-:]\s*)(.*)$", re.I)
SCRIPTURE_RE = re.compile(r"^\s*(?:scripture|verse|text)\s*:\s*(.+)$", re.I)
BULLET_RE = re.compile(r"^\s*(?:[-*•–]|(\d+)[\.\)])\s+(.*)$")

def _norm(s:str)->str: return re.sub(r"\s+"," ",s or "").strip()

def _parse_docx(path: Path) -> List[Dict[str, Any]]:
    doc = Document(str(path))
    studies, cur, outline, points = [], None, [], []
    def flush():
        if not cur: return
        cur["outline"] = _norm(" ".join(outline))
        cur["points"] = [p for p in points if p]
        cur.setdefault("resources", [])
        studies.append(cur.copy())

    for p in doc.paragraphs:
        t = _norm(p.text)
        if not t: continue
        if STUDY_RE.match(t):
            flush(); outline, points = [], []
            m = STUDY_RE.match(t)
            cur = {"title": _norm(m.group(3) or f"Study {m.group(1) or m.group(2)}"),
                   "scripture": "", "outline":"", "points":[], "resources":[]}
            continue
        if not cur: cur = {"title":"Study","scripture":"","outline":"","points":[],"resources":[]}
        if SCRIPTURE_RE.match(t): cur["scripture"] = _norm(SCRIPTURE_RE.match(t).group(1)); continue
        if BULLET_RE.match(t): points.append(_norm(BULLET_RE.match(t).group(2))); continue
        outline.append(t)
    flush(); return studies

def _write_json(data, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w",encoding="utf-8") as f: json.dump(data,f,ensure_ascii=False,indent=2)

def _write_csv(data, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    keys = ["title","scripture","outline","points"]
    with path.open("w",encoding="utf-8",newline="") as f:
        w = csv.DictWriter(f, fieldnames=keys)
        w.writeheader()
        for d in data:
            row = d.copy()
            row["points"] = "; ".join(d.get("points",[]))
            w.writerow(row)

def import_studies(docx: Path = DOCX_PATH) -> List[Dict[str, Any]]:
    """Callable for Flask admin route."""
    data = _parse_docx(docx)
    _write_json(data, JSON_OUT)
    _write_csv(data, CSV_OUT)
    print(f"✅ Imported {len(data)} studies → {JSON_OUT.name} & {CSV_OUT.name}")
    return data

if __name__ == "__main__":
    import_studies()
